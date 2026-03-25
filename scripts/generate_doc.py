"""Generates VPN_KAFKA_BLOCKERS.docx from the markdown content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x1e)
    # Light gray background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


doc = Document()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("VPN + Kafka Connection — Blockers and Fixes", 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph(
    "Every blocker you will likely hit when connecting Google Cloud Run to "
    "AWS MSK Kafka via VPN, and exactly how to resolve each one."
)

# ── What to ask AWS team ──────────────────────────────────────────────────────
doc.add_heading("What to ask your AWS team before starting", 1)
doc.add_paragraph("Get these answers before doing anything else:")
for item in [
    "AWS VPC CIDR where MSK lives (e.g. 10.0.0.0/16)",
    "MSK bootstrap broker string — AWS Console → MSK → cluster → View client information",
    "MSK authentication method: SASL/SCRAM, IAM, or none",
    "SASL username + password (if using SASL/SCRAM)",
    "Kafka topic name(s) to consume",
    "Confirmation that route propagation is enabled on MSK subnet route tables",
    "Confirmation that MSK security group allows TCP 9092/9094 from 10.128.0.0/9",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

# ── Full setup ────────────────────────────────────────────────────────────────
doc.add_heading("Full setup", 1)

# Step 1
doc.add_heading("Step 1 — You: create GCP VPN gateway", 2)
doc.add_paragraph("Fill in your values and run in a terminal with gcloud installed:")
add_code_block(doc, """\
PROJECT_ID="your-gcp-project"
REGION="us-central1"
AWS_VPC_CIDR="10.0.0.0/16"        # get from AWS team
VPC_CONNECTOR_CIDR="10.8.0.0/28"  # any unused /28 in your GCP network

gcloud config set project $PROJECT_ID

# Enable APIs
gcloud services enable compute.googleapis.com vpcaccess.googleapis.com

# Create VPN gateway
gcloud compute vpn-gateways create aws-vpn-gateway \\
  --network=default --region=$REGION

# Create router
gcloud compute routers create vpn-router \\
  --network=default --region=$REGION

# Allow Kafka traffic from AWS VPC
gcloud compute firewall-rules create allow-kafka-from-aws \\
  --network=default --direction=INGRESS --action=ALLOW \\
  --rules=tcp:9092,tcp:9094,tcp:9096 \\
  --source-ranges=$AWS_VPC_CIDR

# Create Serverless VPC Access connector (Cloud Run needs this to use the VPN)
gcloud compute networks vpc-access connectors create kafka-connector \\
  --network=default --region=$REGION \\
  --range=$VPC_CONNECTOR_CIDR \\
  --min-instances=2 --max-instances=3 --machine-type=e2-micro

# Get the external IP — send this to your AWS team
gcloud compute vpn-gateways describe aws-vpn-gateway \\
  --region=$REGION \\
  --format="value(vpnInterfaces[0].ipAddress)"\
""")
doc.add_paragraph("Send the printed IP to your AWS team.")

# Step 2
doc.add_heading("Step 2 — AWS team: create VPN connection", 2)
doc.add_paragraph("Tell your AWS team:")
for item in [
    "Create a Customer Gateway — IP: <GCP external IP from step 1>, Routing: Static",
    "Create a Virtual Private Gateway, attach it to the VPC where MSK lives",
    "Create a Site-to-Site VPN Connection — Virtual Private Gateway: above, Customer Gateway: above, Routing: Static, add route 10.128.0.0/9",
    "Enable route propagation on the route table for MSK subnets: VPC → Route Tables → MSK subnet table → Route Propagation → enable VGW",
    "Update MSK security group — allow inbound TCP 9092 and 9094 from 10.128.0.0/9",
    "Download the VPN config and send back: Tunnel 1 outside IP + pre-shared key, Tunnel 2 outside IP + pre-shared key",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

# Step 3
doc.add_heading("Step 3 — You: create VPN tunnels", 2)
doc.add_paragraph("Fill in the 4 values received from the AWS team:")
add_code_block(doc, """\
AWS_TUNNEL_1_IP="<tunnel 1 outside IP>"
AWS_TUNNEL_1_PSK="<tunnel 1 pre-shared key>"
AWS_TUNNEL_2_IP="<tunnel 2 outside IP>"
AWS_TUNNEL_2_PSK="<tunnel 2 pre-shared key>"

gcloud compute vpn-tunnels create aws-tunnel-1 \\
  --vpn-gateway=aws-vpn-gateway --vpn-gateway-interface=0 \\
  --peer-address=$AWS_TUNNEL_1_IP --shared-secret=$AWS_TUNNEL_1_PSK \\
  --ike-version=2 --region=$REGION --router=vpn-router

gcloud compute vpn-tunnels create aws-tunnel-2 \\
  --vpn-gateway=aws-vpn-gateway --vpn-gateway-interface=1 \\
  --peer-address=$AWS_TUNNEL_2_IP --shared-secret=$AWS_TUNNEL_2_PSK \\
  --ike-version=2 --region=$REGION --router=vpn-router

gcloud compute routers add-interface vpn-router \\
  --interface-name=if-tunnel-1 --vpn-tunnel=aws-tunnel-1 --region=$REGION

gcloud compute routers add-interface vpn-router \\
  --interface-name=if-tunnel-2 --vpn-tunnel=aws-tunnel-2 --region=$REGION

gcloud compute routes create route-to-aws \\
  --network=default \\
  --destination-range=$AWS_VPC_CIDR \\
  --next-hop-vpn-tunnel=aws-tunnel-1 \\
  --next-hop-vpn-tunnel-region=$REGION \\
  --priority=100

# Attach VPC connector to Cloud Run
gcloud run services update kafka-n8n-forwarder \\
  --region=$REGION \\
  --vpc-connector=kafka-connector \\
  --vpc-egress=all-traffic

# Check — both tunnels must show ESTABLISHED
gcloud compute vpn-tunnels list --region=$REGION\
""")
doc.add_paragraph(
    "Both tunnels must show ESTABLISHED before continuing. "
    "If they show FIRST_HANDSHAKE wait 2 minutes and check again."
)

# Step 4
doc.add_heading("Step 4 — You: configure Kafka connection", 2)
doc.add_paragraph("Update PIPELINES_JSON with the MSK broker endpoints and auth, store in Secret Manager, redeploy:")
add_code_block(doc, """\
# PIPELINES_JSON content:
[{
  "name": "prod",
  "kafka_bootstrap_servers": "b-1.cluster.abc.kafka.us-east-1.amazonaws.com:9092",
  "kafka_topics": ["your-topic"],
  "kafka_consumer_group_id": "cloud-run-forwarder-prod",
  "kafka_security_protocol": "SASL_SSL",
  "kafka_sasl_mechanism": "SCRAM-SHA-512",
  "kafka_sasl_username": "your-user",
  "kafka_sasl_password": "your-password",
  "n8n_webhook_url": "https://your-n8n/webhook/xxx",
  "event_type_field": "event_type",
  "event_filters": {}
}]

# Store and redeploy:
echo '[{...}]' | gcloud secrets versions add pipelines-config --data-file=-

gcloud run services update kafka-n8n-forwarder \\
  --region=us-central1 \\
  --update-secrets=PIPELINES_JSON=pipelines-config:latest\
""")

# ── Blockers ──────────────────────────────────────────────────────────────────
doc.add_heading("Blockers", 1)

# Blocker 1
doc.add_heading("Blocker 1 — VPN tunnel stuck at FIRST_HANDSHAKE", 2)
doc.add_paragraph("Tunnels never reach ESTABLISHED.")
add_table(doc,
    ["Cause", "Fix"],
    [
        ["Wrong pre-shared key", "Re-copy PSK from AWS config — no extra spaces or newlines"],
        ["IKE version mismatch", "Try --ike-version=1 instead of 2"],
        ["Wrong Customer Gateway IP", "Must be the GCP VPN gateway external IP, not an internal IP"],
        ["AWS VPN not yet active", "AWS VPN connections take 3–5 min after creation — wait and retry"],
    ]
)
doc.add_paragraph("Check what specifically failed:")
add_code_block(doc, """\
gcloud compute vpn-tunnels describe aws-tunnel-1 --region=us-central1
# Read the "detailedStatus" field\
""")

# Blocker 2
doc.add_heading("Blocker 2 — VPN established but traffic doesn't flow", 2)
doc.add_paragraph(
    "Tunnels show ESTABLISHED but Kafka connections still time out.\n\n"
    "Cause: AWS route table not updated — return traffic from MSK has no path back through the VPN.\n\n"
    "Fix: AWS team must enable route propagation on the MSK subnet route table:\n"
    "AWS Console → VPC → Route Tables → select the table for MSK subnets → Route Propagation tab → enable the Virtual Private Gateway."
)

# Blocker 3
doc.add_heading("Blocker 3 — MSK security group rejects connections", 2)
doc.add_paragraph("VPN works but Kafka connection is refused. AWS team adds inbound rules to the MSK security group:")
add_table(doc,
    ["Protocol", "Port", "Source"],
    [
        ["TCP", "9092", "10.128.0.0/9"],
        ["TCP", "9094", "10.128.0.0/9"],
        ["TCP", "9096", "10.128.0.0/9"],
    ]
)

# Blocker 4
doc.add_heading("Blocker 4 — MSK hostname doesn't resolve from GCP", 2)
doc.add_paragraph(
    "Error: UnknownHostException for broker hostnames like "
    "b-1.cluster.abc.kafka.us-east-1.amazonaws.com\n\n"
    "Why: MSK private DNS is served by Route 53 private hosted zones inside the AWS VPC. "
    "GCP cannot use AWS's internal DNS resolver.\n\n"
    "Fix: Create a Cloud DNS forwarding zone pointing at the AWS VPC DNS resolver. "
    "The AWS DNS resolver IP is always: VPC base CIDR + 2 "
    "(e.g. VPC is 10.0.0.0/16 → resolver is 10.0.0.2)."
)
add_code_block(doc, """\
gcloud dns managed-zones create aws-kafka-dns \\
  --dns-name="kafka.us-east-1.amazonaws.com." \\
  --description="Forward MSK DNS to AWS" \\
  --visibility=private \\
  --networks=default \\
  --forwarding-targets=10.0.0.2   # replace with your VPC base + 2\
""")
doc.add_paragraph("Redeploy Cloud Run after creating the zone.")

# Blocker 5
doc.add_heading("Blocker 5 — Kafka authentication fails", 2)
doc.add_paragraph(
    "Error: SaslAuthenticationException or SSL handshake failed.\n\n"
    "Fix: Match the config to what MSK is actually configured for. "
    "Ask the AWS team which auth method is enabled."
)
doc.add_paragraph("SASL/SCRAM (most common):")
add_code_block(doc, """\
"kafka_security_protocol": "SASL_SSL",
"kafka_sasl_mechanism": "SCRAM-SHA-512",
"kafka_sasl_username": "your-user",
"kafka_sasl_password": "your-password"\
""")
doc.add_paragraph("No auth (plaintext, only safe inside private VPC):")
add_code_block(doc, """\
"kafka_security_protocol": "PLAINTEXT",
"kafka_sasl_mechanism": null,
"kafka_sasl_username": null,
"kafka_sasl_password": null\
""")
doc.add_paragraph("IAM auth: not supported. Ask the AWS team to enable SASL/SCRAM alongside IAM.")

# Blocker 6
doc.add_heading("Blocker 6 — Wrong broker port", 2)
doc.add_paragraph(
    "Connection times out even though VPN and DNS are working. "
    "Get the exact broker string from: AWS Console → MSK → your cluster → View client information."
)
add_table(doc,
    ["Auth method", "Port"],
    [
        ["Plaintext, no auth", "9092"],
        ["TLS, no auth", "9094"],
        ["SASL/SCRAM + TLS", "9096"],
        ["IAM + TLS", "9098"],
    ]
)

# Blocker 7
doc.add_heading("Blocker 7 — Consumer receives no messages", 2)
doc.add_paragraph("Kafka connection succeeds but no events arrive.")
doc.add_paragraph("Cause A — Consumer group conflict:")
doc.add_paragraph(
    "Another consumer (e.g. Lambda) is using the same consumer group ID on the same topic. "
    "Kafka splits partitions between them and Cloud Run may get none."
)
doc.add_paragraph("Fix: use a unique consumer group ID:")
add_code_block(doc, '"kafka_consumer_group_id": "cloud-run-forwarder-prod"')
doc.add_paragraph("Cause B — Wrong offset reset:")
doc.add_paragraph(
    "Consumer group already has committed offsets and auto_offset_reset is 'latest' — it skips old messages."
)
add_code_block(doc, '"kafka_auto_offset_reset": "latest"   // or "earliest" to read from the beginning')

# Blocker 8
doc.add_heading("Blocker 8 — Cloud Run restarts before Kafka connects", 2)
doc.add_paragraph(
    "Health probe fails and Cloud Run keeps restarting the container before the Kafka connection is established.\n\n"
    "Fix: increase failureThreshold in the startup probe in service.yaml to give more time:"
)
add_code_block(doc, """\
startupProbe:
  httpGet:
    path: /ready
  initialDelaySeconds: 5
  periodSeconds: 5
  failureThreshold: 36   # 180 seconds\
""")
doc.add_paragraph("Redeploy after changing:")
add_code_block(doc, "gcloud run services replace service.yaml --region=us-central1")

# ── Diagnostic checklist ──────────────────────────────────────────────────────
doc.add_heading("Diagnostic checklist", 1)
doc.add_paragraph("Run in order to find where the connection is failing:")
add_code_block(doc, """\
# 1. VPN tunnels up?
gcloud compute vpn-tunnels list --region=us-central1

# 2. VPC connector active?
gcloud compute networks vpc-access connectors describe kafka-connector --region=us-central1

# 3. Cloud Run using the connector?
gcloud run services describe kafka-n8n-forwarder --region=us-central1 \\
  --format="value(spec.template.metadata.annotations)"
# Should show: run.googleapis.com/vpc-access-connector: kafka-connector

# 4. App logs — look for "Connecting to Kafka" or error messages
gcloud run services logs read kafka-n8n-forwarder --region=us-central1 --limit=50

# 5. Health check
curl https://your-cloud-run-url/health

# 6. Ready check — fails until Kafka connects
curl https://your-cloud-run-url/ready

# 7. Metrics — messages_consumed goes up when Kafka is working
curl https://your-cloud-run-url/metrics\
""")

# ── Save ──────────────────────────────────────────────────────────────────────
out = "VPN_KAFKA_BLOCKERS.docx"
doc.save(out)
print(f"Saved: {out}")
